#!/usr/bin/env python3
"""
Build a comprehensive Word document cataloging every option/page in the
AddManuChain dashboard, with a screenshot + detailed explanation for each,
followed by a deep-think analysis section.
"""
import os
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

SCREENSHOT_DIR = "/home/z/my-project/screenshots"
OUTPUT = "/home/z/my-project/AddManuChain_Dashboard_Complete_Guide.docx"

# ─── Color palette (emerald/slate/amber — matches the app, no indigo/blue) ───
PRIMARY = RGBColor(0x0F, 0x2A, 0x23)      # deep emerald-slate
ACCENT = RGBColor(0x05, 0x96, 0x69)        # emerald-600
ACCENT_DARK = RGBColor(0x06, 0x5F, 0x46)   # emerald-800
AMBER = RGBColor(0xB4, 0x53, 0x09)         # amber-700
SLATE = RGBColor(0x33, 0x41, 0x55)         # slate-700
LIGHT_BG = "F0FDF4"                         # emerald-50
CARD_BG = "F8FAFC"                          # slate-50
BORDER_CLR = "D1D5DB"

# ─── Page catalog: (section, page_id, title, screenshot, explanation, key_features[]) ───
PAGES = [
    # ── PIPELINE ──
    ("Pipeline", "overview", "Overview / Role Dashboard",
     "01-overview.png",
     "The landing page of the platform. Renders a role-aware dashboard that adapts its KPIs, quick actions, and recent-activity feed to the signed-in user's role (Administrator, OEM Engineer, Print Center Operator, Certification Authority, or End Customer). The admin view shows platform-wide metrics: active orders, printers online, print-center utilization, pending DRM approvals, and a live operations map. It acts as the mission-control landing zone from which every other module is one click away. A role switcher in the sidebar lets an admin impersonate any role to preview their experience without re-logging-in. This is the only page most users see first, so it is designed to surface the 3-5 things that role cares about most within 2 seconds.",
     ["Role-aware widget layout (admin / OEM / operator / authority / customer)",
      "Live KPI cards: orders, printers, utilization, DRM queue, revenue",
      "Recent activity stream (orders, prints, shipments, audit events)",
      "Quick-action shortcuts to the most-used pages for the current role",
      "Operations map showing print centers and active shipments geographically"]),

    ("Pipeline", "orders", "Orders",
     "02-orders.png",
     "The order management pipeline. Every part request flows through here: a customer or operator raises an order for a physical spare part, the system checks whether a certified blueprint exists, routes the order to the nearest qualified print center (or to a traditional supplier if no blueprint exists), and tracks the order through pending → approved → printing → quality-check → shipped → delivered. The page supports filtering by status, priority, part type, and center, plus bulk status updates. Each order row expands to show the full lifecycle: requester, blueprint reference, assigned print center, DRM approval chain, print job ID, and shipment tracking. This is the transactional backbone of the platform — the 'Pipeline' section exists because order flow is the single most important business process.",
     ["Full order lifecycle: pending → approved → printing → QC → shipped → delivered",
      "Filter & search by status, priority, part, center, requester",
      "Inline order detail: blueprint link, DRM chain, print job, shipment",
      "Bulk status updates and CSV export",
      "Auto-routing to the nearest qualified print center based on material + certification"]),

    ("Pipeline", "print_queue", "DRM Approval Queue",
     "03-drm-approval.png",
     "The Digital Rights Management approval pipeline. Before any part can be printed, the OEM that owns the intellectual property must grant a per-print digital license (DRM unlock). This page is the funnel where OEM engineers review incoming print requests against their licensing terms, approve or reject them, and attach conditions (royalty, geographic restriction, quantity cap). The funnel view shows how many requests are pending at each stage: submitted → OEM review → cert-authority review → unlocked → printed. Badges in the sidebar show the live pending count. This stage exists because additive manufacturing makes it trivial to copy a design — DRM is the contractual and technical control that lets OEMs participate in a distributed-print network without losing control of their IP.",
     ["Funnel view of the DRM approval pipeline by stage",
      "Per-request OEM review with approve / reject / condition buttons",
      "Royalty, geographic, and quantity-cap conditions attachable to each unlock",
      "Cert-authority co-sign for regulated parts",
      "Full audit trail of who approved what, when, and under which license terms"]),

    ("Pipeline", "print_approval", "Print Approval",
     "04-print-approval.png",
     "The certification-compliance sign-off queue. After DRM is unlocked but before the printer starts, a certification authority (DNV, Lloyd's Register, ABS, etc.) may need to approve the print for regulated industries (marine, aerospace, defense). This page lists prints awaiting cert sign-off, shows the certification pathway required for each part class, and lets authority users approve, flag, or request re-qualification. It is distinct from DRM Approval because DRM is an IP/commercial control whereas print approval is a safety/regulatory control — a part can pass DRM but still need cert sign-off, or vice versa.",
     ["Cert-pathway-aware approval queue (DNV, LR, ABS, class societies)",
      "Per-print regulatory checklist (material cert, dimensional, NDE)",
      "Approve / flag / request-re-qualification actions",
      "Linkage to the blueprint's certified status and material test reports",
      "Separation of IP approval (DRM) from safety approval (cert)"]),

    ("Pipeline", "emergency", "Emergency Response",
     "05-emergency.png",
     "The fast-path triage page for breakdown scenarios. When a critical part fails on a remote asset (a ship at sea, a rig, a mine), every minute of downtime is expensive. This page collapses the normal order → DRM → cert → print → ship pipeline into a single emergency workflow: the operator enters the broken part, the system instantly surfaces any certified blueprint, the nearest print center with the right material and an open slot, and a one-click 'emergency unlock' that fast-tracks DRM + cert in parallel rather than sequentially. The page shows ETA estimates for each path (print locally vs. print at nearest peer vs. traditional order) so the operator can pick the fastest. This is grounded in the interview insight that remote operations live or die by spare-part lead time.",
     ["Single-screen breakdown → replacement triage",
      "Instant surfacing of certified blueprints for the broken part",
      "Parallel fast-track of DRM + cert approvals (not sequential)",
      "Side-by-side ETA comparison: local print vs. peer print vs. traditional order",
      "Nearest qualified print center with open capacity auto-recommended"]),

    # ── AI INTELLIGENCE ──
    ("AI Intelligence", "ai_agent", "AI Operations Agent",
     "06-ai-agent.png",
     "The autonomous AI operations assistant — the 'next level' AI helper. A user types a natural-language request ('order 5 more thruster bearings', 'which parts are out of stock?', 'generate an onboarding plan based on Robert's knowledge'). The LLM parses the request into a structured plan {intent, tool, parameters, confidence, requiresApproval}. Safe read-only tools (answer question, list low stock, find AM candidates, search knowledge, generate onboarding) auto-execute immediately and return the result. Write tools (create order, adjust inventory, trigger print) return the plan for human approval via a second click. Every action — auto-executed or approved — is written to the AgentActionLog audit trail. A 'Connected Dashboards' bar at the top links to the 6 related dashboards, and each result card shows contextual 'Jump to' buttons that navigate to the most relevant dashboard for follow-up. This is the AI that 'can receive a request and either help or do things automatically' as the user requested.",
     ["Natural-language request parsing into a structured tool plan",
      "8 tools: 5 safe (auto-execute) + 3 write (need approval)",
      "Confidence score + reasoning shown for every plan",
      "Full audit log of every AI action (auto-executed, approved, or failed)",
      "Connected Dashboards cross-link bar + contextual result navigation",
      "Suggested-prompt chips color-coded by tool family"]),

    ("AI Intelligence", "ai_part_scanner", "AI Part Scanner",
     "07-ai-part-scanner.png",
     "The additive-manufacturing candidate scanner. This page runs an AI analysis across the entire physical inventory and scores every part on 6 dimensions (AM suitability, cost impact, lead-time risk, criticality, certification pathway, demand frequency, IP status) to answer: 'which of my spare parts are worth digitizing for 3D printing?'. Each part gets an overall suitability score, a per-dimension breakdown, and an ROI estimate (print cost vs. buy cost vs. lead-time savings). The user can expand any part to see the reasoning, generate a detailed ROI report, and push high-value candidates straight to the Blueprint Library or AM Feasibility page. Grounded in the interview insight that 'not every part is worth digitizing — you need a filter'.",
     ["AI scores every inventory part on 6 AM-suitability dimensions",
      "Overall suitability score + per-dimension radar breakdown",
      "ROI estimate: print cost vs. buy cost vs. lead-time savings",
      "Filterable by score, criticality, demand, IP status",
      "One-click push to Blueprint Library or AM Feasibility for top candidates",
      "AI-generated narrative summarizing the top opportunities"]),

    ("AI Intelligence", "ai_auto_print", "Auto-Print Rules",
     "08-auto-print.png",
     "The smart replenishment rule engine. Instead of a human watching stock levels and manually triggering reprints, this page lets an admin define rules: 'when part X drops below min stock AND a certified blueprint exists AND a print center with the right material is online → auto-trigger a print of quantity Y'. Rules can be scoped by part, site, or material. The page shows a live evaluation of all rules against current inventory, flagging which rules have fired (auto-triggered), which are pending approval (safety threshold), and which are info-only suggestions. Every auto-triggered print is audit-logged. This is the 'when inventory drops, the printer kicks in automatically' feature.",
     ["Rule builder: condition (stock level, blueprint, material, center) → action (auto-print quantity)",
      "Live rule evaluation against current inventory",
      "Three trigger modes: auto_triggered, approval_required, info_only",
      "Per-rule safety threshold (high-value parts need a human confirm)",
      "Full audit trail of every auto-triggered print job"]),

    # ── PEOPLE & KNOWLEDGE ──
    ("People & Knowledge", "workforce_knowledge", "Workforce Knowledge & Memory",
     "09-workforce.png",
     "The institutional-memory platform. Captures the knowledge of retiring senior experts before it walks out the door and transfers it to new hires. Four tabs: (1) Employee Directory — 8 seeded employees (4 retiring seniors + 4 junior hires with mentor links), each with specialties, years, retirement date, and at-risk-of-knowledge-loss flags; (2) Knowledge Library — 10 captured documents (SOPs, troubleshooting guides, procedures, safety bulletins, case studies) authored by seniors, searchable and filterable by category/criticality; (3) Ask a Senior — a RAG chat where a junior types a question and the AI answers grounded ONLY in the captured knowledge docs, citing [KD-xxx] references; (4) Onboarding Generator — pick a retiring senior and the AI generates a 90-day onboarding plan for their replacement drawing on that senior's captured docs (days 1-30/31-60/61-90, must-read docs, must-shadow builds, milestones, risks). Grounded in Jim Granger #23 ('those guys are gone') and Jordan Cumming #10 ('knowledge capture = highest immediate value').",
     ["Employee directory with mentor/mentee links and retirement risk flags",
      "Knowledge Library: 10 captured SOPs, procedures, troubleshooting guides",
      "Ask-a-Senior RAG chat — grounded answers citing [KD-xxx] document IDs",
      "90-day Onboarding Plan Generator driven by a retiring senior's captured knowledge",
      "Connected Dashboards bar + Ask-a-Senior → AI Agent handoff button"]),

    # ── RESOURCES ──
    ("Resources", "blueprints", "Blueprint Library",
     "10-blueprints.png",
     "The certified CAD blueprint repository. Every part that can be 3D-printed needs a certified blueprint — a CAD file that has been qualified for additive manufacturing (print parameters validated, material specified, dimensional tolerance proven, NDE pass). This page lists all 16 seeded blueprints with their part number, material, certification status, owning OEM, royalty terms, and print count. Blueprints can be filtered by material, certification status, and OEM. Each blueprint row links to its print history, material test reports, and the orders that consumed it. This is the 'digital inventory' that makes on-demand printing possible — without a certified blueprint, a part cannot be printed.",
     ["16 certified CAD blueprints with material, tolerances, print parameters",
      "Certification status: certified / pending / re-qualification",
      "OEM ownership + royalty terms per blueprint",
      "Print history and material test report linkage",
      "Filter by material, cert status, OEM, part class"]),

    ("Resources", "centers", "Print Centers",
     "11-centers.png",
     "The certified AM facility directory. Lists all 5 print centers (Atlantic XL, DNV Calgary, LR Montreal, St. John's AM, Victoria Marine) with their location, materials supported, machine list, capacity, utilization, and certification status. Each center has a status (online / offline / maintenance) and a real-time queue depth. The page lets an admin see at a glance which centers can print which materials and how busy they are — critical for order routing and emergency response. A center detail view shows its printer fleet, current jobs, and upcoming schedule.",
     ["5 certified print centers with location, materials, capacity",
      "Real-time status (online/offline/maintenance) and queue depth",
      "Per-center printer fleet and utilization metrics",
      "Material capability matrix (which center can print which alloy/polymer)",
      "Certification status per center (DNV, LR, ABS approved)"]),

    ("Resources", "my_printers", "My Printers",
     "12-my-printers.png",
     "The on-site printer management page. For an operator who has their own 3D printers on-site, this page lists their fleet with status, current job, material loaded, maintenance schedule, and print history. It supports scheduling (queue jobs for a specific printer), material loading tracking, and maintenance alerts. This is the operator's day-to-day view of their own hardware, distinct from the platform-wide Print Centers page.",
     ["On-site printer fleet with live status and current job",
      "Material loaded + remaining filament/powder/resin tracking",
      "Job scheduling and queue management per printer",
      "Maintenance schedule and alerts",
      "Print history per printer with success/failure rate"]),

    ("Resources", "peer_printers", "Peer Printers",
     "13-peer-printers.png",
     "The peer-to-peer printer network — an 'Airbnb for 3D printers'. Independent operators and small shops can list their certified printers on the platform, and customers in need of a part can rent capacity on the nearest qualified peer printer. This page lists available peer printers with location, materials, certification, hourly rate, and availability calendar. It dramatically expands the platform's print capacity without AddManuChain owning the hardware. Grounded in the cooperative-manufacturing interview theme.",
     ["Peer-to-peer marketplace of certified 3D printers",
      "Per-printer location, materials, certification, hourly rate",
      "Availability calendar and booking system",
      "Rating and review system for peer operators",
      "Geo-search: find the nearest qualified peer printer"]),

    ("Resources", "shipments", "Shipments",
     "14-shipments.png",
     "The logistics tracking page. Once a print completes and passes QC, the finished part must ship to the customer. This page tracks every shipment: origin center, destination, carrier, tracking number, ETA, and status (packed → in-transit → delivered). It supports filtering by status, date, and destination, and integrates with carrier tracking APIs. For emergency orders, the page flags expedited shipments and shows live ETA.",
     ["Full shipment lifecycle: packed → in-transit → delivered",
      "Carrier integration with tracking numbers and live ETA",
      "Filter by status, date range, origin, destination",
      "Expedited/emergency shipment flagging",
      "Proof-of-delivery capture and history"]),

    ("Resources", "materials", "Materials",
     "15-materials.png",
     "The raw-material inventory page. Tracks the filament, resin, and metal powder stock at each print center and warehouse: material type, alloy/grade, quantity on hand, lot/batch number, supplier, expiry, and certification. This is critical because AM materials have shelf lives and lot-level traceability requirements — a print is only as certified as the material lot it was printed from. The page flags materials nearing expiry or below safety stock.",
     ["Raw material stock per center/warehouse (filament, resin, powder)",
      "Lot/batch traceability with supplier and certification",
      "Expiry tracking with near-expiry alerts",
      "Safety-stock thresholds with reorder alerts",
      "Material-to-blueprint compatibility matrix"]),

    ("Resources", "lab_portal", "Lab & Testing Portal",
     "16-lab-portal.png",
     "The testing and qualification portal. New parts and materials must be tested (tensile, fatigue, dimensional, NDE) before they can be certified for production use. This page manages test requests, equipment scheduling, and certification reports. An OEM or operator submits a test request, the lab schedules it on the right equipment, runs the test, and uploads the certified report. The report then feeds back into the blueprint's certification status. This closes the loop between 'print a part' and 'prove it meets spec'.",
     ["Test request submission (tensile, fatigue, dimensional, NDE)",
      "Lab equipment scheduling and availability",
      "Certified test report upload and storage",
      "Linkage to blueprint certification status",
      "Test history per part and per material lot"]),

    ("Resources", "physical_inventory", "Physical Inventory",
     "17-physical-inventory.png",
     "The traditional spare-parts warehouse view. Lists all 14 physical parts across all sites with part number, name, site, quantity, min stock, condition, unit cost, last used, last inspected, and notes. This is the 'before AM' baseline — the parts you already own sitting on shelves. It supports manual stock adjustments, condition updates, and transaction history (received, consumed, condemned, inspected). The Smart Inventory Console builds on top of this data to add AI analysis and dual-mode control.",
     ["14 physical parts across all sites with full stock metadata",
      "Per-part condition (new / used / condemned), unit cost, last used/inspected",
      "Transaction history: received, consumed, condemned, inspected",
      "Manual stock adjustments with reason codes",
      "Multi-site warehouse view with site filter"]),

    ("Resources", "smart_inventory", "Smart Inventory Console",
     "18-smart-inventory.png",
     "The dual-mode inventory management console — manual OR AI. The flagship 'inventory management that can be handled manually or by AI' feature. Two toggle modes share the same parts table and audit trail: Manual Mode lets an operator adjust stock by hand with quick +1/-1 buttons and a full adjust dialog (new quantity, reason, performed-by). AI Assist Mode runs an LLM analysis across the entire inventory and returns prioritized decisions (reorder, digitize_for_am, transfer_surplus, flag_slow_mover, condemn, safety_stock_adjust) — each with confidence, est. impact, and reasoning. The user can approve/reject individual decisions or hit 'Auto-Execute All High-Confidence' to let the AI handle everything above 0.8 confidence. Every action — AI or manual — is written to the InventoryAiDecision audit trail. A 'Connected Dashboards' bar links to 6 related dashboards; decision cards show contextual 'Open Blueprint Library' / 'Open Orders' buttons. Grounded in Cameron Munro #81.",
     ["Dual-mode toggle: Manual (operator) vs AI Assist (LLM-driven)",
      "6 KPI cards: total parts, out of stock, below min, healthy, total value, condemned",
      "AI analysis returns 6 decision types with confidence + est. impact",
      "One-click approve/reject per decision or batch auto-execute ≥ 0.8 confidence",
      "Unified audit trail — every AI and manual action logged to InventoryAiDecision",
      "Connected Dashboards bar + contextual decision-card navigation"]),

    ("Resources", "digital_inventory", "Digital Inventory",
     "19-digital-inventory.png",
     "The AI-driven digital spare-parts forecast. Rather than tracking what you physically own, this page predicts what you WILL need and whether you should hold it physically or digitally. It forecasts demand per part, identifies parts that should be digitized (hold the blueprint, print on demand) vs. physically stocked, and calculates the working-capital savings of a digital-first strategy. This is the strategic complement to Physical Inventory — together they answer 'what should I own vs. what should I be able to print on demand?'.",
     ["AI demand forecast per part (next 30/60/90 days)",
      "Digital-vs-physical holding recommendation per part",
      "Working-capital savings calculation for digital-first strategy",
      "Lead-time risk analysis (parts where print-on-demand beats buy-and-hold)",
      "Conversion roadmap: which physical parts to digitize first"]),

    # ── PARTNERS & IP ──
    ("Partners & IP", "partners", "OEM Partners",
     "20-partners.png",
     "The OEM partner relationship directory. Lists every OEM that has licensed blueprints to the platform, with their contact, licensing terms, royalty rate, blueprint count, and print volume. This page manages the commercial relationship — contracts, royalty statements, partner health score. It is the account-management view for the AddManuChain team's partner managers.",
     ["OEM partner directory with licensing terms and royalty rate",
      "Per-partner blueprint count and print volume",
      "Royalty statement generation and history",
      "Partner health score (active blueprints, approval turnaround, dispute rate)",
      "Contract and contact management"]),

    ("Partners & IP", "oem_portal", "OEM Self-Service Portal",
     "21-oem-portal.png",
     "The OEM's self-service view. An OEM partner logs in and sees only their own blueprints, their own incoming DRM approval queue, their royalty earnings, and their print history. They can upload new blueprints, set licensing terms (geographic, quantity, royalty), approve/reject DRM requests, and pull royalty reports. This page exists so OEMs don't need AddManuChain staff to manage their IP day-to-day — it's the SaaS self-service layer that makes the partner model scale.",
     ["OEM-scoped view: only their blueprints, DRM queue, royalties",
      "Blueprint upload with licensing terms (geographic, quantity, royalty)",
      "Self-service DRM approve/reject with condition attachment",
      "Royalty earnings dashboard and statement download",
      "Print history and consumer-anonymized usage analytics"]),

    ("Partners & IP", "ip_library", "IP Library",
     "22-ip-library.png",
     "The intellectual-property rights management page. Tracks every IP license in the platform: which OEM owns which blueprint, what license terms apply (exclusive, non-exclusive, geographic, temporal), royalty per print, and the full royalty payment history. This is the legal/financial system of record for IP — distinct from the Blueprint Library (which is the technical CAD file) and OEM Partners (which is the relationship). It answers 'who owes whom how much for which print'.",
     ["IP license registry per blueprint (exclusive/non-exclusive, geo, temporal)",
      "Royalty rate per print and cumulative royalty owed/paid",
      "License-term enforcement (geographic + quantity caps)",
      "Royalty payment history and reconciliation",
      "IP dispute tracking and resolution"]),

    ("Partners & IP", "cooperative", "Digital Cooperative",
     "23-cooperative.png",
     "The member cooperative blueprint pool. Members of the AddManuChain cooperative can contribute certified blueprints to a shared pool and draw any part from the pool to print at any member facility. This page browses the cooperative pool, shows contribution/consumption balance per member, and manages the royalty-sharing rules among members. It is the 'open-source' layer on top of the commercial OEM-IP layer — members get access to a broader library in exchange for contributing their own designs. Grounded in the digital-cooperative interview theme.",
     ["Shared cooperative blueprint pool browsable by all members",
      "Per-member contribution / consumption balance",
      "Royalty-sharing rules among cooperative members",
      "Member-only print access at any member facility",
      "Contribution quality bar (certified + tested) for pool inclusion"]),

    # ── COMPLIANCE ──
    ("Compliance", "certifications", "Certifications",
     "24-certifications.png",
     "The certification registry. Lists every certification held by the platform, its centers, and its parts: certifying body (DNV, LR, ABS), cert type (material, process, part, facility), scope, issue/expiry date, and status. Certifications are the regulatory license to operate — without the right cert, a printed part cannot be used in a regulated application. This page is the single source of truth for 'what are we certified to do, and when does it expire'.",
     ["Certification registry: material, process, part, facility certs",
      "Certifying body tracking (DNV, LR, ABS, class societies)",
      "Issue/expiry dates with renewal alerts",
      "Scope and conditions per certification",
      "Status: active / pending / expired / suspended"]),

    ("Compliance", "authorities", "Certification Authorities",
     "25-authorities.png",
     "The certifying-body relationship page. Manages AddManuChain's relationships with DNV, Lloyd's Register, ABS, and other class societies: contact, account manager, active certifications, audit history, and pending requests. This is the account-management view for compliance — analogous to OEM Partners but for regulators instead of IP owners. It tracks audit cycles, renewal windows, and the authority's approval throughput.",
     ["Certifying-body relationship directory (DNV, LR, ABS)",
     "Per-authority contact and account manager",
     "Active certifications and audit history per authority",
     "Renewal window tracking and audit cycle management",
     "Authority approval throughput metrics"]),

    ("Compliance", "audit", "Audit Chain",
     "26-audit.png",
     "The immutable system audit log. Every state-changing action across the entire platform — order created, DRM unlocked, print started, stock adjusted, AI agent executed, user login, role change — is written to the audit chain. This page is the compliance and forensic view: filterable by actor, action type, entity, and timestamp. It supports CSV export for regulatory submissions and legal hold. The audit chain is the platform's defense in any dispute — 'who did what, when, and from where'.",
     ["Immutable audit log of every state-changing action platform-wide",
      "Filter by actor, action type, entity, timestamp",
      "CSV export for regulatory submission",
      "Legal-hold capability for dispute preservation",
      "Tamper-evident chaining of audit entries"]),

    # ── INTELLIGENCE ──
    ("Intelligence", "analytics", "Analytics",
     "27-analytics.png",
     "The operational analytics dashboard. Standard business-intelligence charts: orders over time, print success rate, center utilization, revenue, top parts, geographic distribution. This is the 'how is the business doing' view for executives and account managers. It is the baseline analytics layer that the Advanced Analytics page extends with ML.",
     ["Orders, revenue, print volume over time",
      "Print success rate and quality metrics",
      "Per-center utilization and throughput",
      "Top parts by volume and by revenue",
      "Geographic distribution of orders and prints"]),

    ("Intelligence", "advanced_analytics", "Advanced Analytics",
     "28-advanced-analytics.png",
     "The ML-powered predictive analytics layer. Builds on the baseline Analytics page with predictive models: demand forecasting (which parts will be needed next month), failure prediction (which printers are likely to fail), quality prediction (which prints are at risk of defect), and cost optimization (which parts should shift from buy to print). This is where the platform's data assets turn into forward-looking decisions.",
     ["ML demand forecasting per part (30/60/90 day horizon)",
      "Printer failure prediction with maintenance recommendation",
      "Print quality risk scoring (pre-print defect probability)",
      "Cost optimization: buy-vs-print recommendation per part",
      "Model confidence intervals and accuracy backtesting"]),

    ("Intelligence", "sc_intelligence", "Supply Chain Intelligence",
     "29-sc-intelligence.png",
     "The supply-chain sovereignty and optimization dashboard. Three lenses: (1) Sovereignty — what % of critical parts are domestically printable vs. import-dependent, with risk scoring; (2) LEAN analysis — waste identification across the parts pipeline (over-production, waiting, excess inventory); (3) Pareto working-capital optimizer — the 80/20 of which parts tie up the most capital and whether digitizing them would release it. Grounded in the supply-chain-sovereignty interview theme that remote operations cannot afford import dependence.",
     ["Sovereignty dashboard: domestic-print % vs. import dependence",
      "LEAN waste analysis across the parts pipeline",
      "Pareto working-capital optimizer (top capital-tying parts)",
      "Risk scoring for single-source and import-dependent parts",
      "Digitization roadmap ranked by capital-release impact"]),

    ("Intelligence", "material_properties", "Material Properties Library",
     "30-material-properties.png",
     "The tested AM material mechanical-data library — the MMPDS equivalent for additive manufacturing. Every material qualified on the platform has a datasheet: tensile strength, yield, elongation, fatigue, hardness, thermal properties, print parameters, and the lot-level test data backing each value. This is the engineering reference that lets an OEM or authority trust that a printed part will perform as designed. Without it, AM parts cannot be certified for safety-critical use.",
     ["Tested mechanical properties per AM material (tensile, fatigue, hardness)",
      "Print parameters (layer height, speed, temp, inert gas)",
      "Lot-level test data traceability",
      "MMPDS-equivalent statistical basis (A/B/S basis)",
      "Material-to-application suitability scoring"]),

    ("Intelligence", "feasibility", "AM Feasibility Triage",
     "31-feasibility.png",
     "The 30-second AM feasibility verdict. An engineer enters a part (name, geometry, material, quantity, urgency, industry, function) and the AI returns an instant go/no-go verdict on whether the part is suitable for additive manufacturing, with a confidence score, the top 3 reasons for/against, and a recommended next step (digitize, redesign for AM, or buy traditionally). This is the front-door triage that prevents wasted effort digitizing parts that will never print well.",
     ["30-second AI go/no-go AM feasibility verdict",
      "Inputs: geometry, material, quantity, urgency, industry, function",
      "Confidence score + top 3 reasons for/against",
      "Recommended next step: digitize / redesign / buy",
      "Connected Dashboards link to Part Scanner + Blueprint Library"]),

    ("Intelligence", "customer_success", "Customer Success",
     "32-customer-success.png",
     "The customer-engagement and success-management page. Tracks every customer's onboarding progress, training completion, parts-printed history, satisfaction score, and renewal/expansion risk. The customer-success team uses this to run QBRs (quarterly business reviews), flag at-risk accounts, and identify expansion opportunities. It is the post-sale account-management view.",
     ["Customer onboarding progress and training completion",
      "Per-customer parts-printed history and satisfaction score",
      "Renewal and expansion-risk flagging",
      "QBR-ready account summaries",
      "Expansion-opportunity identification"]),

    # ── SYSTEM ──
    ("System", "reports", "Reports & Export",
     "33-reports.png",
     "The report-generation and export hub. Lets any user assemble a report (orders, prints, inventory, audit, financial) as PDF, Excel, or CSV with date range and filter controls. Scheduled reports can be emailed on a cadence. This is the 'get data out of the system for a meeting or a regulator' page.",
     ["Report builder: orders, prints, inventory, audit, financial",
      "Export formats: PDF, Excel, CSV",
      "Date-range and filter controls per report",
      "Scheduled email delivery (daily/weekly/monthly)",
      "Saved report templates per user"]),

    ("System", "notifications", "Notifications",
     "34-notifications.png",
     "The real-time alert and notification center. In-app + email + (optional) SMS alerts for: order status changes, DRM approvals, low-stock thresholds, print completions, failures, audit events, and system health. Each user configures their notification preferences per channel. The page shows the live feed and lets the user mark-read, snooze, or jump to the source entity.",
     ["Real-time in-app + email + SMS alert center",
      "Alert types: order, DRM, stock, print, failure, audit, system",
      "Per-user channel preferences (in-app/email/SMS)",
      "Mark-read, snooze, jump-to-source actions",
      "Alert digest configuration (instant / hourly / daily)"]),

    ("System", "search", "Advanced Search",
     "35-search.png",
     "The full-text advanced search across every entity in the platform. One search bar that queries orders, blueprints, parts, shipments, users, audit logs, knowledge docs — with faceted filters (entity type, date, status, owner). This is the 'I know it exists somewhere, find it' tool that complements the role-specific dashboards.",
     ["Unified full-text search across all entities",
      "Faceted filters: entity type, date, status, owner",
      "Saved searches and search history",
      "Jump-to-record from search results",
      "Search-as-you-type with result previews"]),

    ("System", "batch", "Batch Operations",
     "36-batch.png",
     "The bulk-action workspace. Lets an admin select multiple entities (orders, prints, inventory items) and apply a single action to all of them: status update, reassign, export, delete, tag. This is the power-user tool for managing large datasets efficiently — e.g. 'mark all pending orders from last week as reviewed'.",
     ["Multi-select across orders, prints, inventory, users",
      "Bulk actions: status update, reassign, export, delete, tag",
      "Filter-then-select-all workflow",
      "Batch preview before commit",
      "Batch audit-log entry per bulk action"]),

    ("System", "monitoring", "System Monitoring",
     "37-monitoring.png",
     "The platform health dashboard. Real-time metrics: API response times, error rates, database query performance, printer connectivity, background-job queue depth, and uptime. Alerts fire when SLAs breach. This is the SRE / ops view that keeps the platform running.",
     ["Real-time API response time and error-rate metrics",
      "Database query performance and slow-query log",
      "Printer connectivity and offline-alert monitoring",
      "Background-job queue depth and failure rate",
      "Uptime SLO tracking and breach alerting"]),

    ("System", "integrations", "Integrations Hub",
     "38-integrations.png",
     "The third-party integration manager. Connects the platform to external systems: ERP (SAP, Oracle), PLM (Siemens Teamcenter), CAD (SolidWorks, Fusion 360), carrier APIs (FedEx, UPS), and identity providers (Okta, Azure AD). Each integration has a config page, connection health, and sync history.",
     ["ERP integrations (SAP, Oracle, NetSuite)",
      "PLM/CAD integrations (Teamcenter, SolidWorks, Fusion 360)",
      "Carrier APIs (FedEx, UPS, DHL)",
      "Identity providers (Okta, Azure AD, SAML)",
      "Per-integration health, sync history, and field mapping"]),

    ("System", "automation", "Workflow Automation",
     "39-automation.png",
     "The workflow-automation rule manager. Define event-triggered automations: 'when an order is created → auto-assign to nearest center', 'when stock drops below min → create a reorder', 'when a print fails → notify the operator + log a ticket'. Rules are built via a visual trigger-condition-action editor. This is the no-code automation layer that reduces manual handoffs.",
     ["Visual trigger-condition-action rule builder",
      "Event triggers: order created, stock low, print done, audit event",
      "Actions: assign, notify, create-order, adjust-stock, call-webhook",
      "Rule simulation and test-run before activation",
      "Per-rule execution log and error handling"]),

    ("System", "workflow_builder", "Workflow Builder",
     "40-workflow-builder.png",
     "The advanced visual workflow canvas. A step-by-step flow designer where complex multi-step processes (new-part onboarding, emergency response, cert renewal) are modeled as a flowchart of tasks, decisions, and handoffs. Each step can be automated or routed to a human. This is the process-engineering tool that sits above the simpler Automation rules.",
     ["Drag-and-drop visual workflow canvas",
      "Step types: task, decision, parallel-branch, wait, webhook",
      "Per-step assignee or automation binding",
      "Workflow versioning and publish/draft",
      "Live execution instance tracking per workflow"]),

    ("System", "users", "User Management",
     "41-users.png",
     "The team and permissions manager. Add/remove users, assign roles (admin, OEM engineer, print center operator, cert authority, end customer), scope their access (which centers, which blueprints, which orders), and reset credentials. This is the IAM (identity and access management) layer for the platform.",
     ["User CRUD with role assignment (5 roles)",
      "Per-user access scoping (centers, blueprints, orders)",
      "Role-based permission matrix",
      "Password reset and MFA enforcement",
      "User activity and last-login audit"]),

    ("System", "api", "API Management",
     "42-api.png",
     "The API key and developer-access manager. Issue/revoke API keys, set per-key rate limits and scopes, view usage analytics, and access the developer docs. This is the page a customer's integration team uses to connect their ERP or build a custom app on top of the platform.",
     ["API key issuance and revocation",
      "Per-key rate limits and permission scopes",
      "Usage analytics (calls, errors, latency)",
      "Developer documentation and sandbox",
      "Webhook endpoint registration and delivery log"]),

    ("System", "dashboards", "Custom Dashboards",
     "43-dashboards.png",
     "The personalized dashboard builder. Each user can assemble their own dashboard from a widget library (KPI cards, charts, tables, maps) — pin the metrics they care about and hide the rest. Saved dashboards become the user's default landing view. This is the personalization layer on top of the role-based Overview.",
     ["Widget library: KPI cards, charts, tables, maps, feeds",
      "Drag-and-drop dashboard layout editor",
      "Per-user saved dashboards (multiple per user)",
      "Set a saved dashboard as the default landing view",
      "Share dashboards with team members"]),

    ("System", "mobile_dashboard", "Mobile Dashboard",
     "44-mobile.png",
     "The mobile-optimized dashboard view. Renders the platform in a phone-friendly layout: stacked cards, bottom tab bar, touch-optimized controls, and offline-capable critical views (e.g. a field technician can view an assigned print job even without signal). This is the field-ops companion to the desktop dashboard.",
     ["Mobile-optimized card layout (no horizontal scroll)",
      "Bottom tab bar navigation (Home / Queue / Inventory / Alerts / Me)",
      "Touch-optimized controls (44px+ tap targets)",
      "Offline-capable critical views (assigned jobs, part lookup)",
      "Push-notification integration"]),

    ("System", "one-click-ordering", "One-Click AI Order Automation",
     "45-one-click.png",
     "The AI-assisted rapid-order page. A single input where a user describes what they need in plain language ('I need a thruster bearing housing for the Atlantic XL by Friday') and the AI parses it, finds the matching blueprint, checks stock, routes to the nearest center, pre-fills the order form, and submits it in one click. This is the consumer-grade '1-click' experience layered on top of the enterprise order pipeline.",
     ["Natural-language order entry ('I need a X for Y by Z')",
      "AI parses → matches blueprint → checks stock → routes center",
      "Pre-filled order form review before 1-click submit",
      "Auto-applies the fastest path (print vs. transfer vs. buy)",
      "Order confirmation with ETA and tracking link"]),

    ("System", "settings", "Settings",
     "46-settings.png",
     "The account and platform settings page. Per-user preferences (theme, language, timezone, notification channels) and per-admin platform config (organization profile, billing, branding, feature flags, environment variables). This is the config layer that everything else reads from.",
     ["Per-user preferences: theme, language, timezone, notifications",
      "Organization profile and branding (logo, colors, domain)",
      "Billing and subscription plan management",
      "Feature-flag toggles per environment",
      "Environment variable and integration-key vault"]),
]

# ─── Build the document ───────────────────────────────────────────────────
doc = Document()

# Page setup — A4 landscape would be too wide; use portrait A4
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def add_page_break(p=None):
    if p is None:
        p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

# ─── COVER PAGE ───
# Title
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AddManuChain Platform")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = ACCENT_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Complete Dashboard Guide")
run.font.size = Pt(28)
run.font.color.rgb = PRIMARY

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Every Option, Every Page, Every Feature — Explained")
run.font.size = Pt(14)
run.font.italic = True
run.font.color.rgb = SLATE

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A detailed walkthrough of all 46 dashboard pages across 9 sections,\nwith screenshots and in-depth explanations of each feature.")
run.font.size = Pt(12)
run.font.color.rgb = SLATE

for _ in range(8):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Manufacturing Supply Chain Management Platform  |  Next.js 16 + AI")
run.font.size = Pt(10)
run.font.color.rgb = ACCENT

add_page_break()

# ─── TABLE OF CONTENTS (manual) ───
p = doc.add_paragraph()
run = p.add_run("Table of Contents")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = PRIMARY
p.space_after = Pt(18)

doc.add_paragraph()

# Group by section
from collections import OrderedDict
sections = OrderedDict()
for sec, pid, title, img, expl, feats in PAGES:
    sections.setdefault(sec, []).append((pid, title))

toc_num = 1
for sec_name, pages in sections.items():
    p = doc.add_paragraph()
    run = p.add_run(sec_name)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = ACCENT_DARK
    p.space_after = Pt(4)
    for pid, title in pages:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(title)
        run.font.size = Pt(11)
        run.font.color.rgb = SLATE
        p.paragraph_format.left_indent = Cm(1)
    doc.add_paragraph()

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Deep-Think Analysis")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = ACCENT_DARK
p = doc.add_paragraph()
run = p.add_run("A strategic assessment of the platform's architecture, strengths, gaps, and future direction.")
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = SLATE

add_page_break()

# ─── PAGES ───
page_num = 1
current_section = None
for sec, pid, title, img, expl, feats in PAGES:
    # Section divider
    if sec != current_section:
        if current_section is not None:
            add_page_break()
        current_section = sec
        p = doc.add_paragraph()
        run = p.add_run(f"Section: {sec}")
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = ACCENT
        p.paragraph_format.space_after = Pt(6)

        # Section intro line
        section_intros = {
            "Pipeline": "The core transactional flow — from order request through DRM and certification approval to emergency response.",
            "AI Intelligence": "The AI layer that makes the platform 'next-level' — an autonomous agent, a part scanner, and auto-print rules.",
            "People & Knowledge": "Institutional-memory capture and transfer — keeping retiring experts' knowledge in the platform.",
            "Resources": "The physical and digital assets — blueprints, print centers, printers, shipments, materials, inventory (physical, smart, and digital).",
            "Partners & IP": "The commercial and IP layer — OEM partners, the OEM self-service portal, IP rights, and the member cooperative.",
            "Compliance": "The regulatory layer — certifications, certifying-authority relationships, and the immutable audit chain.",
            "Intelligence": "The analytics and engineering layer — operational analytics, ML predictions, supply-chain intelligence, material data, AM feasibility, and customer success.",
            "System": "The platform-administration layer — reports, notifications, search, batch ops, monitoring, integrations, automation, workflow builder, users, API, custom dashboards, mobile, one-click ordering, and settings.",
        }
        p = doc.add_paragraph()
        run = p.add_run(section_intros.get(sec, ""))
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = SLATE
        p.paragraph_format.space_after = Pt(14)

    # Page heading
    p = doc.add_heading(level=2)
    run = p.add_run(f"{page_num}. {title}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    # Meta line
    p = doc.add_paragraph()
    run = p.add_run(f"Section: {sec}   |   Page ID: {pid}")
    run.font.size = Pt(9)
    run.font.color.rgb = SLATE
    run.font.italic = True
    p.paragraph_format.space_after = Pt(8)

    # Screenshot
    img_path = os.path.join(SCREENSHOT_DIR, img)
    if os.path.exists(img_path):
        # Get aspect ratio
        im = Image.open(img_path)
        w, h = im.size
        # Display width 16cm, preserve aspect ratio
        disp_w = Cm(16)
        disp_h = int(disp_w * h / w)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=disp_w)
        p.paragraph_format.space_after = Pt(4)

        # Caption
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Figure {page_num}: {title} dashboard")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = SLATE
        p.paragraph_format.space_after = Pt(10)
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[Screenshot not found: {img}]")
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # Explanation heading
    p = doc.add_paragraph()
    run = p.add_run("What this page does")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = ACCENT_DARK
    p.paragraph_format.space_after = Pt(4)

    # Explanation body
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(expl)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(10)

    # Key features heading
    p = doc.add_paragraph()
    run = p.add_run("Key features")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = ACCENT_DARK
    p.paragraph_format.space_after = Pt(4)

    # Key features bullets
    for feat in feats:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(feat)
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()  # spacing
    page_num += 1

# ─── DEEP-THINK ANALYSIS ───
add_page_break()

p = doc.add_heading(level=1)
run = p.add_run("Deep-Think Analysis")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = PRIMARY
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
run = p.add_run("A strategic assessment of the AddManuChain platform — its architecture, strengths, latent gaps, and the trajectory it implies.")
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = SLATE
p.paragraph_format.space_after = Pt(18)

# Analysis sections
analyses = [
    ("1. The Platform's Core Thesis",
     "AddManuChain is not a 3D-printing tool. It is a supply-chain sovereignty platform that uses additive manufacturing as its execution layer. The thesis is: remote operations (ships, rigs, mines, defense bases) are crippled by spare-part lead time, and the combination of (a) certified digital blueprints, (b) a distributed network of certified print centers, (c) DRM-protected IP licensing, and (d) AI-driven inventory decisions collapses that lead time from weeks to hours. Every page in this 46-page catalog is in service of that thesis. The Pipeline section exists because order flow is the transactional proof. The AI Intelligence section exists because human judgement cannot keep up with the optimization surface. The People & Knowledge section exists because the senior experts who know which parts matter are retiring. The Compliance section exists because regulated industries cannot use uncertified parts. The System section exists because an enterprise platform must be operable. The catalog is not a feature list — it is a complete operating model for distributed, certified, on-demand manufacturing."),

    ("2. Architectural Strengths",
     "Three architectural decisions stand out. First, the separation of DRM Approval (IP/commercial control) from Print Approval (safety/regulatory control) is correct — these are independent concerns that naive platforms conflate, which either blocks legitimate prints or rubber-stamps unsafe ones. Second, the dual-mode Smart Inventory Console (manual OR AI) is the right human-in-the-loop pattern: it gives operators a familiar manual path while the AI path earns trust through an audit trail, rather than forcing an all-or-nothing AI adoption. Third, the Connected Dashboards cross-navigation layer (added in Phase 5) turns 46 isolated pages into a coherent workflow — an AI Agent result can jump to Smart Inventory, which can jump to Workforce Knowledge, which can hand back to the AI Agent. The platform's value compounds when these connections are traversable, not just present."),

    ("3. The AI Layer's Maturity Curve",
     "The platform's AI features sit at three maturity tiers. Tier 1 (informational): the AI Part Scanner, AM Feasibility Triage, and Advanced Analytics — these produce recommendations that a human acts on elsewhere. Tier 2 (assistive): the Smart Inventory Console's AI Assist Mode and the Auto-Print Rules — these propose actions with confidence scores and require human approval for high-stakes execution. Tier 3 (autonomous): the AI Operations Agent — this auto-executes safe read-only tools and surfaces write tools for approval. The progression from Tier 1 to Tier 3 is the right order: earn trust with read-only AI, then assistive AI, then autonomous AI. The next tier (Tier 4: fully autonomous write execution with guardrails) is technically buildable today — the AgentActionLog and InventoryAiDecision audit trails already exist — but should wait until the audit data shows that Tier 3's approval-reject rate is low enough that auto-execution would not cause regret."),

    ("4. The Latent Gap: Closed-Loop Learning",
     "The platform logs every action (audit chain, agent action log, inventory AI decisions) but does not yet close the loop on learning. When the AI suggests 'digitize part X for AM' and the human approves, the print either succeeds or fails — but that outcome is not yet fed back into the AI's confidence model. Similarly, when Ask-a-Senior cites KD-007 and the user finds the answer unhelpful, that signal is not yet used to re-rank the knowledge base. The single highest-leverage next feature is a feedback loop: every AI output gets a thumb-up/thumb-down, every executed decision gets an outcome label (success / failure / partial), and those labels retrain the confidence models. Without this, the AI's confidence scores are static heuristics; with it, they become empirical. This is the difference between an AI that looks smart and an AI that gets smart."),

    ("5. The People Risk Is Bigger Than the Tech Risk",
     "The Workforce Knowledge page is the most strategically important page in the platform, and it is the least technologically flashy. The interview grounding is explicit: 'critical assembly procedures existed only in the heads of experienced technicians who had retired or passed away.' Every other page in the platform assumes the knowledge of which part to print, from which blueprint, in which material, to which tolerance — and that knowledge lives in senior experts who are retiring within 5 years. The 90-day Onboarding Plan Generator and Ask-a-Senior RAG are good starts, but the platform should treat knowledge capture as a first-class operational KPI: every senior's knowledge-doc count, every junior's onboarding-plan completion, every Ask-a-Senior answer's helpfulness score. The day a senior retires with zero captured docs should be a platform-level alert, not a silent loss."),

    ("6. The Cooperative vs. Commercial Tension",
     "The platform hosts two business models in tension. The commercial model (OEM Partners, OEM Portal, IP Library) is IP-protected, royalty-bearing, and controlled. The cooperative model (Digital Cooperative, Peer Printers) is shared, member-governed, and open. Both are valid, but they cannot be silently mixed — a member who contributes a blueprint to the cooperative pool while the same blueprint is under OEM license creates a rights conflict. The platform currently keeps them in separate nav sections, which is correct, but it should add an explicit rights-conflict detector: when a blueprint is about to be contributed to the cooperative, the system should verify no active OEM license covers the same part. This is a small feature with large legal consequence."),

    ("7. The Mobile and Field Layer Is Underserved",
     "The Mobile Dashboard page exists, but the platform's value proposition — spare parts for remote operations — implies a field technician on a rig or a ship who needs to: identify a broken part by photo, check if a blueprint exists, trigger an emergency print, and track the shipment. That workflow is only partially served by the current mobile view (which is a responsive re-skin of the desktop dashboard). A purpose-built field-technician mobile app (photo-to-part VLM identification, offline blueprint cache, emergency-print one-tap, shipment ETA push) would be the highest-NPS feature for the actual end users. The VLM and web-search skills available in the platform's AI stack are exactly the tools needed to build the photo-to-part identifier."),

    ("8. Recommendation: The Next Three Features",
     "If I had to pick the next three features to build, in priority order: (1) Closed-loop AI feedback — every AI output gets an outcome label that retrains confidence; this makes every other AI feature better simultaneously. (2) Field-technician mobile companion — a purpose-built mobile app with photo-to-part VLM identification and offline-cached blueprints; this reaches the actual end user the platform is built for. (3) Certification Pathway Builder — grounded in the Saxty #45 interview's 7-step qualification framework; this turns the current 'certifications are a registry' into 'certifications are a guided workflow,' which is the difference between a compliance tool and a compliance partner. All three are buildable on the existing architecture without schema migrations, and all three compound the value of the 46 pages already shipped."),

    ("9. Final Verdict",
     "AddManuChain is a serious platform. The 46-page catalog is not feature bloat — each page maps to a real stakeholder need, and the interview grounding (Cumming, Granger, Dalpe, Munro, Saxty, Kobalch) shows the design is evidence-led rather than assumption-led. The AI layer is at the right maturity (assistive, not autonomous-by-default). The audit trails are in place for the day autonomy is earned. The cross-dashboard navigation turns the catalog into a workflow. The gaps (closed-loop learning, field mobile, certification pathway builder) are clear and buildable. The platform's biggest risk is not technical — it is the people risk of retiring seniors, which the Workforce Knowledge page addresses but cannot fully solve. The platform's biggest opportunity is closing the AI feedback loop, which would make every AI feature better simultaneously. The trajectory is correct; the next 90 days should be feedback loop, field mobile, and certification pathway, in that order."),
]

for heading, body in analyses:
    p = doc.add_heading(level=2)
    run = p.add_run(heading)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = ACCENT_DARK
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(body)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(10)

# ─── SAVE ───
doc.save(OUTPUT)
print(f"✓ Document saved: {OUTPUT}")
print(f"  Pages documented: {len(PAGES)}")
print(f"  Sections: {len(sections)}")
print(f"  Deep-think analysis sections: {len(analyses)}")

# File size
size = os.path.getsize(OUTPUT)
print(f"  File size: {size / 1024 / 1024:.2f} MB ({size:,} bytes)")
